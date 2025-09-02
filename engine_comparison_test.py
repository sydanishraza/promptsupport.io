#!/usr/bin/env python3
"""
Training Engine vs Knowledge Engine Comparison Test
Tests both engines to understand capabilities and recommend enhancement focus
"""

import os
import sys
import json
import asyncio
import aiohttp
import aiofiles
from datetime import datetime
import tempfile
import requests
from docx import Document
from docx.shared import Inches
import io
import base64

# Get backend URL from environment
BACKEND_URL = os.getenv('REACT_APP_BACKEND_URL', 'https://knowledge-engine-7.preview.emergentagent.com')
API_BASE = f"{BACKEND_URL}/api"

class EngineComparisonTest:
    def __init__(self):
        self.session = None
        self.test_results = {
            'knowledge_engine': {
                'content_upload': {'status': 'pending', 'details': []},
                'processing_quality': {'status': 'pending', 'details': []},
                'asset_library': {'status': 'pending', 'details': []}
            },
            'training_engine': {
                'templates': {'status': 'pending', 'details': []},
                'sessions': {'status': 'pending', 'details': []},
                'processing': {'status': 'pending', 'details': []},
                'image_handling': {'status': 'pending', 'details': []}
            }
        }
        
    async def setup_session(self):
        """Setup HTTP session for testing"""
        connector = aiohttp.TCPConnector(ssl=False)
        timeout = aiohttp.ClientTimeout(total=300)  # 5 minute timeout
        self.session = aiohttp.ClientSession(
            connector=connector,
            timeout=timeout,
            headers={
                'User-Agent': 'EngineComparison-Test/1.0',
                'Accept': 'application/json'
            }
        )
        
    async def cleanup_session(self):
        """Cleanup HTTP session"""
        if self.session:
            await self.session.close()
    
    def create_test_docx_with_images(self, filename: str) -> str:
        """Create a test DOCX file with embedded images for testing"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Engine Comparison Test Document', 0)
            
            # Add introduction
            intro = doc.add_paragraph(
                'This document tests both Knowledge Engine and Training Engine capabilities. '
                'It contains structured content with multiple sections to evaluate '
                'content processing, image handling, and article generation quality.'
            )
            
            # Add first section with content
            doc.add_heading('Section 1: Content Processing Capabilities', level=1)
            doc.add_paragraph(
                'Both engines should be able to extract and process this content effectively. '
                'The Knowledge Engine focuses on content chunking and library integration, '
                'while the Training Engine emphasizes article generation and image contextual placement.'
            )
            
            # Add second section
            doc.add_heading('Section 2: Image Processing and Media Assets', level=1)
            doc.add_paragraph(
                'This section tests image extraction and processing capabilities. '
                'The Training Engine has advanced HTML preprocessing pipeline for accurate image reinsertion, '
                'while the Knowledge Engine focuses on asset library integration and media management.'
            )
            
            # Add third section
            doc.add_heading('Section 3: Article Generation and Quality', level=1)
            doc.add_paragraph(
                'The final comparison focuses on output quality and structure. '
                'Training Engine generates comprehensive articles with contextual image placement, '
                'while Knowledge Engine creates structured content chunks for knowledge base integration.'
            )
            
            # Save document
            doc.save(filename)
            print(f"✅ Created comparison test DOCX file: {filename}")
            return filename
            
        except Exception as e:
            print(f"❌ Failed to create test DOCX: {e}")
            return None
    
    async def test_knowledge_engine(self):
        """Test Knowledge Engine capabilities"""
        print("\n🔍 TESTING KNOWLEDGE ENGINE")
        
        try:
            # Test 1: Content Upload
            test_file = os.path.join(tempfile.gettempdir(), 'engine_comparison_test.docx')
            docx_path = self.create_test_docx_with_images(test_file)
            
            if docx_path and os.path.exists(docx_path):
                upload_url = f"{API_BASE}/content/upload"
                
                with open(docx_path, 'rb') as f:
                    file_data = f.read()
                
                data = aiohttp.FormData()
                data.add_field('file', file_data, filename='engine_comparison_test.docx', 
                              content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                data.add_field('metadata', json.dumps({'test': 'engine_comparison', 'source': 'knowledge_engine_test'}))
                
                print(f"📤 Testing Knowledge Engine upload: {len(file_data)} bytes")
                
                async with self.session.post(upload_url, data=data) as response:
                    if response.status == 200:
                        result = await response.json()
                        
                        if result.get('status') == 'completed':
                            self.test_results['knowledge_engine']['content_upload']['status'] = 'passed'
                            self.test_results['knowledge_engine']['content_upload']['details'].append(
                                f"✅ Upload successful: {result.get('extracted_content_length', 0)} chars, {result.get('chunks_created', 0)} chunks"
                            )
                            
                            # Store for further testing
                            self.knowledge_engine_result = result
                        else:
                            self.test_results['knowledge_engine']['content_upload']['status'] = 'failed'
                            self.test_results['knowledge_engine']['content_upload']['details'].append(
                                f"❌ Upload failed: {result.get('status', 'unknown')}"
                            )
                    else:
                        self.test_results['knowledge_engine']['content_upload']['status'] = 'failed'
                        self.test_results['knowledge_engine']['content_upload']['details'].append(
                            f"❌ HTTP {response.status}: {await response.text()}"
                        )
                
                # Cleanup
                try:
                    os.remove(docx_path)
                except:
                    pass
            
            # Test 2: Asset Library
            assets_url = f"{API_BASE}/assets"
            async with self.session.get(assets_url) as response:
                if response.status == 200:
                    assets_data = await response.json()
                    assets_count = len(assets_data.get('assets', []))
                    
                    self.test_results['knowledge_engine']['asset_library']['status'] = 'passed'
                    self.test_results['knowledge_engine']['asset_library']['details'].append(
                        f"✅ Asset Library accessible: {assets_count} assets"
                    )
                    
                    # Check for recent assets
                    recent_assets = [asset for asset in assets_data.get('assets', []) 
                                   if asset.get('source') == 'docx_extraction']
                    
                    if recent_assets:
                        self.test_results['knowledge_engine']['asset_library']['details'].append(
                            f"📷 Found {len(recent_assets)} DOCX-extracted images"
                        )
                else:
                    self.test_results['knowledge_engine']['asset_library']['status'] = 'failed'
                    self.test_results['knowledge_engine']['asset_library']['details'].append(
                        f"❌ Asset Library not accessible: HTTP {response.status}"
                    )
            
            # Test 3: Content Library
            content_library_url = f"{API_BASE}/content-library"
            async with self.session.get(content_library_url) as response:
                if response.status == 200:
                    library_data = await response.json()
                    articles_count = len(library_data.get('articles', []))
                    
                    self.test_results['knowledge_engine']['processing_quality']['status'] = 'passed'
                    self.test_results['knowledge_engine']['processing_quality']['details'].append(
                        f"✅ Content Library accessible: {articles_count} articles"
                    )
                else:
                    self.test_results['knowledge_engine']['processing_quality']['status'] = 'failed'
                    self.test_results['knowledge_engine']['processing_quality']['details'].append(
                        f"❌ Content Library not accessible: HTTP {response.status}"
                    )
                    
        except Exception as e:
            print(f"❌ Knowledge Engine test failed: {e}")
    
    async def test_training_engine(self):
        """Test Training Engine capabilities"""
        print("\n🔍 TESTING TRAINING ENGINE")
        
        try:
            # Test 1: Templates endpoint
            templates_url = f"{API_BASE}/training/templates"
            async with self.session.get(templates_url) as response:
                if response.status == 200:
                    templates_data = await response.json()
                    templates_count = len(templates_data.get('templates', []))
                    
                    self.test_results['training_engine']['templates']['status'] = 'passed'
                    self.test_results['training_engine']['templates']['details'].append(
                        f"✅ Templates endpoint accessible: {templates_count} templates"
                    )
                else:
                    self.test_results['training_engine']['templates']['status'] = 'failed'
                    self.test_results['training_engine']['templates']['details'].append(
                        f"❌ Templates endpoint not accessible: HTTP {response.status}"
                    )
            
            # Test 2: Sessions endpoint
            sessions_url = f"{API_BASE}/training/sessions"
            async with self.session.get(sessions_url) as response:
                if response.status == 200:
                    sessions_data = await response.json()
                    sessions_count = len(sessions_data.get('sessions', []))
                    
                    self.test_results['training_engine']['sessions']['status'] = 'passed'
                    self.test_results['training_engine']['sessions']['details'].append(
                        f"✅ Sessions endpoint accessible: {sessions_count} sessions"
                    )
                else:
                    self.test_results['training_engine']['sessions']['status'] = 'failed'
                    self.test_results['training_engine']['sessions']['details'].append(
                        f"❌ Sessions endpoint not accessible: HTTP {response.status}"
                    )
            
            # Test 3: Processing endpoint with DOCX
            test_file = os.path.join(tempfile.gettempdir(), 'training_engine_test.docx')
            docx_path = self.create_test_docx_with_images(test_file)
            
            if docx_path and os.path.exists(docx_path):
                process_url = f"{API_BASE}/training/process"
                
                with open(docx_path, 'rb') as f:
                    file_data = f.read()
                
                data = aiohttp.FormData()
                data.add_field('file', file_data, filename='training_engine_test.docx', 
                              content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                data.add_field('template_id', 'document_upload_processing')
                
                print(f"📤 Testing Training Engine processing: {len(file_data)} bytes")
                
                async with self.session.post(process_url, data=data) as response:
                    if response.status == 200:
                        result = await response.json()
                        
                        if result.get('success'):
                            articles = result.get('articles', [])
                            images_processed = result.get('images_processed', 0)
                            
                            self.test_results['training_engine']['processing']['status'] = 'passed'
                            self.test_results['training_engine']['processing']['details'].append(
                                f"✅ Processing successful: {len(articles)} articles generated"
                            )
                            
                            self.test_results['training_engine']['image_handling']['status'] = 'passed' if images_processed > 0 else 'failed'
                            self.test_results['training_engine']['image_handling']['details'].append(
                                f"{'✅' if images_processed > 0 else '❌'} Images processed: {images_processed}"
                            )
                            
                            # Analyze article quality
                            for i, article in enumerate(articles[:3]):  # Check first 3 articles
                                if isinstance(article, dict):
                                    title = article.get('title', 'Untitled')
                                    content = article.get('content', '')
                                    word_count = len(content.split()) if content else 0
                                    has_images = '<img' in content or '<figure' in content
                                    
                                    self.test_results['training_engine']['processing']['details'].append(
                                        f"📄 Article {i+1}: '{title[:40]}...' ({word_count} words, {'with' if has_images else 'no'} images)"
                                    )
                        else:
                            self.test_results['training_engine']['processing']['status'] = 'failed'
                            self.test_results['training_engine']['processing']['details'].append(
                                f"❌ Processing failed: {result.get('error', 'unknown error')}"
                            )
                    else:
                        response_text = await response.text()
                        self.test_results['training_engine']['processing']['status'] = 'failed'
                        self.test_results['training_engine']['processing']['details'].append(
                            f"❌ HTTP {response.status}: {response_text[:200]}..."
                        )
                
                # Cleanup
                try:
                    os.remove(docx_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"❌ Training Engine test failed: {e}")
    
    async def run_comparison_tests(self):
        """Run comprehensive comparison tests"""
        print("🚀 Starting Engine Comparison Testing Suite")
        print(f"🔗 Backend URL: {BACKEND_URL}")
        
        await self.setup_session()
        
        try:
            await self.test_knowledge_engine()
            await self.test_training_engine()
            
        finally:
            await self.cleanup_session()
        
        # Print comparison results
        self.print_comparison_summary()
    
    def print_comparison_summary(self):
        """Print comprehensive comparison results"""
        print("\n" + "="*80)
        print("🎯 ENGINE COMPARISON TEST RESULTS")
        print("="*80)
        
        # Knowledge Engine Results
        print("\n📚 KNOWLEDGE ENGINE RESULTS:")
        ke_passed = 0
        ke_total = 0
        
        for test_name, result in self.test_results['knowledge_engine'].items():
            ke_total += 1
            status = result['status']
            
            if status == 'passed':
                status_icon = "✅"
                ke_passed += 1
            elif status == 'failed':
                status_icon = "❌"
            else:
                status_icon = "⏳"
            
            print(f"  {status_icon} {test_name.replace('_', ' ').title()}: {status.upper()}")
            for detail in result['details']:
                print(f"     • {detail}")
        
        # Training Engine Results
        print("\n🏋️ TRAINING ENGINE RESULTS:")
        te_passed = 0
        te_total = 0
        
        for test_name, result in self.test_results['training_engine'].items():
            te_total += 1
            status = result['status']
            
            if status == 'passed':
                status_icon = "✅"
                te_passed += 1
            elif status == 'failed':
                status_icon = "❌"
            else:
                status_icon = "⏳"
            
            print(f"  {status_icon} {test_name.replace('_', ' ').title()}: {status.upper()}")
            for detail in result['details']:
                print(f"     • {detail}")
        
        # Overall Comparison
        print(f"\n📊 COMPARISON SUMMARY:")
        print(f"   Knowledge Engine: {ke_passed}/{ke_total} tests passed ({ke_passed/ke_total*100:.1f}%)")
        print(f"   Training Engine:  {te_passed}/{te_total} tests passed ({te_passed/te_total*100:.1f}%)")
        
        # Recommendations
        print(f"\n💡 RECOMMENDATIONS:")
        
        if ke_passed >= te_passed:
            print("   🎯 FOCUS ON KNOWLEDGE ENGINE ENHANCEMENT:")
            print("     • Knowledge Engine shows better stability and functionality")
            print("     • Content upload and processing working reliably")
            print("     • Asset library and content library integration operational")
            print("     • Recommend enhancing Knowledge Engine with:")
            print("       - Improved image processing and contextual placement")
            print("       - Enhanced content coverage and article generation")
            print("       - Better media asset management")
        else:
            print("   🎯 FOCUS ON TRAINING ENGINE ENHANCEMENT:")
            print("     • Training Engine shows superior capabilities")
            print("     • Advanced article generation and image processing")
            print("     • HTML preprocessing pipeline for accurate image reinsertion")
            print("     • Recommend enhancing Training Engine with:")
            print("       - Improved stability and error handling")
            print("       - Better content library integration")
            print("       - Enhanced processing performance")
        
        print("="*80)

async def main():
    """Main test execution"""
    tester = EngineComparisonTest()
    await tester.run_comparison_tests()

if __name__ == "__main__":
    asyncio.run(main())