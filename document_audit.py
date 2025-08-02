#!/usr/bin/env python3
"""
Document Audit Script - Master Product Management Guide
Analyzes the original DOCX document to compare with Knowledge Engine output
"""

import os
import sys
from docx import Document
from PIL import Image
import io
import requests
import json
import time
from collections import defaultdict

def analyze_docx_document(docx_path):
    """Analyze the original DOCX document structure"""
    print(f"🔍 Analyzing original document: {docx_path}")
    
    try:
        doc = Document(docx_path)
        
        # Document metadata
        print(f"\n📋 DOCUMENT METADATA:")
        print(f"Title: {doc.core_properties.title or 'N/A'}")
        print(f"Author: {doc.core_properties.author or 'N/A'}")
        print(f"Subject: {doc.core_properties.subject or 'N/A'}")
        
        # Count document elements
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        
        # Analyze headings and structure
        headings = defaultdict(list)
        total_text_length = 0
        paragraph_styles = defaultdict(int)
        
        print(f"\n📊 DOCUMENT STRUCTURE:")
        
        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name
                paragraph_styles[style_name] += 1
                total_text_length += len(para.text)
                
                # Identify headings
                if 'Heading' in style_name or style_name == 'Title':
                    level = 1
                    if 'Heading 1' in style_name or style_name == 'Title':
                        level = 1
                    elif 'Heading 2' in style_name:
                        level = 2
                    elif 'Heading 3' in style_name:
                        level = 3
                    elif 'Heading 4' in style_name:
                        level = 4
                    
                    headings[level].append(para.text.strip())
        
        print(f"Total paragraphs: {total_paragraphs}")
        print(f"Total tables: {total_tables}")
        print(f"Total text length: {total_text_length:,} characters")
        
        print(f"\n📝 PARAGRAPH STYLES:")
        for style, count in sorted(paragraph_styles.items()):
            print(f"  {style}: {count}")
        
        print(f"\n🏗️ HEADING STRUCTURE:")
        total_headings = 0
        for level in sorted(headings.keys()):
            print(f"  Level {level} headings: {len(headings[level])}")
            total_headings += len(headings[level])
            for heading in headings[level][:5]:  # Show first 5 headings
                print(f"    - {heading[:60]}{'...' if len(heading) > 60 else ''}")
            if len(headings[level]) > 5:
                print(f"    ... and {len(headings[level]) - 5} more")
        
        print(f"Total headings: {total_headings}")
        
        # Analyze embedded images
        print(f"\n🖼️ EMBEDDED IMAGES ANALYSIS:")
        image_count = 0
        image_sizes = []
        
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    image_part = rel.target_part
                    image_data = image_part.blob
                    image_sizes.append(len(image_data))
                    
                    # Try to get image dimensions
                    try:
                        img = Image.open(io.BytesIO(image_data))
                        width, height = img.size
                        print(f"  Image {image_count}: {len(image_data):,} bytes, {width}x{height}px, format: {img.format}")
                    except Exception as e:
                        print(f"  Image {image_count}: {len(image_data):,} bytes, format: {image_part.content_type}")
        
        except Exception as e:
            print(f"  Error analyzing images: {e}")
        
        print(f"Total images found: {image_count}")
        if image_sizes:
            print(f"Average image size: {sum(image_sizes)//len(image_sizes):,} bytes")
            print(f"Largest image: {max(image_sizes):,} bytes")
            print(f"Smallest image: {min(image_sizes):,} bytes")
        
        # Expected article split analysis
        print(f"\n📑 EXPECTED ARTICLE SPLIT ANALYSIS:")
        chapter_headings = headings.get(1, []) + headings.get(2, [])
        print(f"Major sections (H1+H2): {len(chapter_headings)}")
        print(f"Expected articles (based on structure): {max(6, len(chapter_headings))}")
        
        # Content density analysis
        avg_chars_per_paragraph = total_text_length / max(1, total_paragraphs - total_headings)
        print(f"Average content per paragraph: {avg_chars_per_paragraph:.0f} characters")
        
        return {
            'total_paragraphs': total_paragraphs,
            'total_tables': total_tables,
            'total_text_length': total_text_length,
            'total_headings': total_headings,
            'total_images': image_count,
            'headings_by_level': dict(headings),
            'paragraph_styles': dict(paragraph_styles),
            'expected_articles': max(6, len(chapter_headings)),
            'image_sizes': image_sizes
        }
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        return None

def test_knowledge_engine_processing(docx_path):
    """Test the document through the Knowledge Engine"""
    print(f"\n🧪 TESTING KNOWLEDGE ENGINE PROCESSING")
    print("=" * 60)
    
    backend_url = "https://c14dc277-70df-425b-a9d5-f1d91d1168d4.preview.emergentagent.com/api"
    
    try:
        # Upload the document
        print("📤 Uploading document to Knowledge Engine...")
        
        with open(docx_path, 'rb') as f:
            files = {'file': ('Master_Product_Management_Guide.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'metadata': '{}'}
            
            response = requests.post(f'{backend_url}/content/upload', files=files, data=data, timeout=120)
            
        if response.status_code == 200:
            print("✅ Upload successful!")
            result = response.json()
            print(f"Response: {result.get('message', 'Success')}")
            
            # Wait for processing
            print("⏳ Waiting for processing...")
            time.sleep(10)
            
            # Get generated articles
            print("📥 Retrieving generated articles...")
            lib_response = requests.get(f'{backend_url}/content-library', timeout=30)
            
            if lib_response.status_code == 200:
                articles_data = lib_response.json()
                all_articles = articles_data.get('articles', [])
                
                # Find articles from this document
                product_management_articles = [
                    a for a in all_articles 
                    if ('product' in a.get('title', '').lower() and 'management' in a.get('title', '').lower()) 
                    or 'master' in a.get('title', '').lower()
                    or a.get('metadata', {}).get('original_filename', '').startswith('Master_Product_Management_Guide')
                ]
                
                print(f"\n📊 KNOWLEDGE ENGINE RESULTS:")
                print(f"Total articles generated: {len(product_management_articles)}")
                
                total_generated_content = 0
                total_images_embedded = 0
                html_tags_count = 0
                markdown_patterns_count = 0
                
                for i, article in enumerate(product_management_articles, 1):
                    title = article.get('title', 'Untitled')
                    content = article.get('content', '')
                    summary = article.get('summary', 'No summary')
                    
                    # Count content
                    content_length = len(content)
                    total_generated_content += content_length
                    
                    # Count HTML vs Markdown
                    html_tags = content.count('<') + content.count('>')
                    markdown_patterns = content.count('##') + content.count('**') + content.count('- ')
                    
                    html_tags_count += html_tags
                    markdown_patterns_count += markdown_patterns
                    
                    # Count images
                    images_in_article = content.count('<img')
                    total_images_embedded += images_in_article
                    
                    print(f"\n  📄 Article {i}: {title[:60]}{'...' if len(title) > 60 else ''}")
                    print(f"     Summary: {summary[:80]}{'...' if len(summary) > 80 else ''}")
                    print(f"     Content length: {content_length:,} characters")
                    print(f"     Images embedded: {images_in_article}")
                    print(f"     HTML tags: {html_tags}, Markdown patterns: {markdown_patterns}")
                    print(f"     AI processed: {article.get('metadata', {}).get('ai_processed', False)}")
                    print(f"     Model used: {article.get('metadata', {}).get('ai_model', 'Unknown')}")
                
                # Overall analysis
                print(f"\n📈 PROCESSING ANALYSIS:")
                print(f"Total content generated: {total_generated_content:,} characters")
                print(f"Total images embedded: {total_images_embedded}")
                print(f"HTML tags total: {html_tags_count}")
                print(f"Markdown patterns total: {markdown_patterns_count}")
                print(f"HTML dominance: {html_tags_count/(html_tags_count+markdown_patterns_count)*100:.1f}%")
                
                return {
                    'articles_generated': len(product_management_articles),
                    'total_content_length': total_generated_content,
                    'total_images_embedded': total_images_embedded,
                    'html_tags': html_tags_count,
                    'markdown_patterns': markdown_patterns_count,
                    'articles': product_management_articles
                }
                
            else:
                print(f"❌ Failed to retrieve content library: {lib_response.status_code}")
                print(f"Error: {lib_response.text}")
                return None
                
        else:
            print(f"❌ Upload failed: {response.status_code}")
            print(f"Error: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Error testing Knowledge Engine: {e}")
        return None

def compare_original_vs_generated(original_analysis, generated_analysis):
    """Compare original document with generated articles"""
    print(f"\n📊 COMPARISON: ORIGINAL vs GENERATED")
    print("=" * 60)
    
    if not original_analysis or not generated_analysis:
        print("❌ Cannot compare - missing analysis data")
        return
    
    print(f"📄 CONTENT COVERAGE:")
    original_chars = original_analysis['total_text_length']
    generated_chars = generated_analysis['total_content_length']
    coverage_ratio = (generated_chars / original_chars) * 100 if original_chars > 0 else 0
    
    print(f"  Original document: {original_chars:,} characters")
    print(f"  Generated articles: {generated_chars:,} characters")
    print(f"  Coverage ratio: {coverage_ratio:.1f}%")
    
    if coverage_ratio < 80:
        print("  ⚠️  LOW COVERAGE - Significant content may be missing")
    elif coverage_ratio > 120:
        print("  ✅ HIGH COVERAGE - Content expanded with AI enhancements")
    else:
        print("  ✅ GOOD COVERAGE - Most content preserved")
    
    print(f"\n🖼️ IMAGE PROCESSING:")
    original_images = original_analysis['total_images']
    generated_images = generated_analysis['total_images_embedded']
    image_ratio = (generated_images / original_images) * 100 if original_images > 0 else 0
    
    print(f"  Original images: {original_images}")
    print(f"  Embedded in articles: {generated_images}")
    print(f"  Embedding ratio: {image_ratio:.1f}%")
    
    if image_ratio < 50:
        print("  ❌ POOR IMAGE EMBEDDING - Most images missing from articles")
    elif image_ratio < 80:
        print("  ⚠️  PARTIAL IMAGE EMBEDDING - Some images missing")
    else:
        print("  ✅ GOOD IMAGE EMBEDDING - Most images preserved")
    
    print(f"\n📑 ARTICLE STRUCTURE:")
    expected_articles = original_analysis['expected_articles']
    actual_articles = generated_analysis['articles_generated']
    
    print(f"  Expected articles: {expected_articles}")
    print(f"  Generated articles: {actual_articles}")
    
    if actual_articles < expected_articles * 0.5:
        print("  ❌ UNDER-SPLITTING - Too few articles created")
    elif actual_articles > expected_articles * 1.5:
        print("  ⚠️  OVER-SPLITTING - Too many articles created")
    else:
        print("  ✅ GOOD SPLITTING - Appropriate number of articles")
    
    print(f"\n🎨 FORMATTING QUALITY:")
    html_dominance = generated_analysis['html_tags'] / (generated_analysis['html_tags'] + generated_analysis['markdown_patterns']) * 100
    print(f"  HTML dominance: {html_dominance:.1f}%")
    
    if html_dominance > 70:
        print("  ✅ EXCELLENT FORMATTING - Proper HTML output")
    elif html_dominance > 50:
        print("  ✅ GOOD FORMATTING - Mostly HTML with some Markdown")
    else:
        print("  ❌ POOR FORMATTING - Too much Markdown, insufficient HTML")

if __name__ == "__main__":
    docx_path = "/app/Master_Product_Management_Guide.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ Document not found: {docx_path}")
        sys.exit(1)
    
    print("🔍 MASTER PRODUCT MANAGEMENT GUIDE - DOCUMENT AUDIT")
    print("=" * 70)
    
    # Analyze original document
    original_analysis = analyze_docx_document(docx_path)
    
    # Test Knowledge Engine processing
    generated_analysis = test_knowledge_engine_processing(docx_path)
    
    # Compare results
    compare_original_vs_generated(original_analysis, generated_analysis)
    
    print(f"\n✅ AUDIT COMPLETE")