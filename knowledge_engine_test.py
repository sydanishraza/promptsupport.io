#!/usr/bin/env python3
"""
Knowledge Engine Testing Suite
Tests the Knowledge Engine's content upload, processing, and image handling capabilities
"""

import os
import sys
import json
import asyncio
import aiohttp
import aiofiles
from datetime import datetime
import tempfile
import requests
from docx import Document
from docx.shared import Inches
import io
import base64

# Get backend URL from environment
BACKEND_URL = os.getenv('REACT_APP_BACKEND_URL', 'https://c14dc277-70df-425b-a9d5-f1d91d1168d4.preview.emergentagent.com')
API_BASE = f"{BACKEND_URL}/api"

class KnowledgeEngineTest:
    def __init__(self):
        self.session = None
        self.test_results = {
            'content_upload': {'status': 'pending', 'details': []},
            'image_processing': {'status': 'pending', 'details': []},
            'content_coverage': {'status': 'pending', 'details': []},
            'article_generation': {'status': 'pending', 'details': []},
            'cors_issues': {'status': 'pending', 'details': []}
        }
        
    async def setup_session(self):
        """Setup HTTP session for testing"""
        connector = aiohttp.TCPConnector(ssl=False)
        timeout = aiohttp.ClientTimeout(total=300)  # 5 minute timeout
        self.session = aiohttp.ClientSession(
            connector=connector,
            timeout=timeout,
            headers={
                'User-Agent': 'KnowledgeEngine-Test/1.0',
                'Accept': 'application/json'
            }
        )
        
    async def cleanup_session(self):
        """Cleanup HTTP session"""
        if self.session:
            await self.session.close()
    
    def create_test_docx_with_images(self, filename: str) -> str:
        """Create a test DOCX file with embedded images for testing"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Knowledge Engine Test Document', 0)
            
            # Add introduction
            intro = doc.add_paragraph(
                'This is a comprehensive test document for the Knowledge Engine. '
                'It contains multiple sections with embedded images to test the '
                'image processing and contextual placement capabilities.'
            )
            
            # Add first section with content
            doc.add_heading('Section 1: Content Processing Overview', level=1)
            doc.add_paragraph(
                'The Knowledge Engine is designed to extract and process content from various document formats. '
                'This section tests the basic content extraction capabilities including text formatting, '
                'structure preservation, and comprehensive coverage of document elements.'
            )
            
            # Create a simple test image (1x1 pixel PNG)
            test_image_data = base64.b64decode(
                'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChAGA'
                'VaUNAAAAAElFTkSuQmCC'
            )
            
            # Save test image temporarily
            temp_img_path = os.path.join(tempfile.gettempdir(), 'test_image.png')
            with open(temp_img_path, 'wb') as f:
                f.write(test_image_data)
            
            # Add image to document
            try:
                doc.add_picture(temp_img_path, width=Inches(2))
                doc.add_paragraph('Figure 1: Test image for Knowledge Engine processing')
            except Exception as img_error:
                print(f"⚠️ Could not add image to DOCX: {img_error}")
            
            # Add second section
            doc.add_heading('Section 2: Image Processing Capabilities', level=1)
            doc.add_paragraph(
                'This section focuses on testing the image extraction and contextual placement features. '
                'The Knowledge Engine should be able to extract embedded images, save them to the asset library, '
                'and place them contextually within generated articles.'
            )
            
            # Add another image
            try:
                doc.add_picture(temp_img_path, width=Inches(1.5))
                doc.add_paragraph('Figure 2: Second test image for contextual placement testing')
            except Exception as img_error:
                print(f"⚠️ Could not add second image to DOCX: {img_error}")
            
            # Add third section
            doc.add_heading('Section 3: Article Generation Quality', level=1)
            doc.add_paragraph(
                'The final test focuses on the quality of generated articles. This includes checking for '
                'comprehensive content coverage, proper HTML structure, professional formatting, and '
                'accurate image placement within the generated content.'
            )
            
            # Save document
            doc.save(filename)
            
            # Cleanup temp image
            try:
                os.remove(temp_img_path)
            except:
                pass
                
            print(f"✅ Created test DOCX file: {filename}")
            return filename
            
        except Exception as e:
            print(f"❌ Failed to create test DOCX: {e}")
            return None
    
    async def test_content_upload_api(self):
        """Test 1: Content Upload API with DOCX file"""
        print("\n🔍 TEST 1: Content Upload API")
        
        try:
            # Create test DOCX file
            test_file = os.path.join(tempfile.gettempdir(), 'knowledge_engine_test.docx')
            docx_path = self.create_test_docx_with_images(test_file)
            
            if not docx_path or not os.path.exists(docx_path):
                self.test_results['content_upload']['status'] = 'failed'
                self.test_results['content_upload']['details'].append('Failed to create test DOCX file')
                return
            
            # Test file upload
            upload_url = f"{API_BASE}/content/upload"
            
            # Prepare multipart form data
            with open(docx_path, 'rb') as f:
                file_data = f.read()
            
            # Create form data
            data = aiohttp.FormData()
            data.add_field('file', file_data, filename='knowledge_engine_test.docx', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            data.add_field('metadata', json.dumps({'test': True, 'source': 'knowledge_engine_test'}))
            
            print(f"📤 Uploading test DOCX to: {upload_url}")
            print(f"📊 File size: {len(file_data)} bytes")
            
            async with self.session.post(upload_url, data=data) as response:
                response_text = await response.text()
                
                print(f"📥 Response status: {response.status}")
                print(f"📥 Response headers: {dict(response.headers)}")
                
                if response.status == 200:
                    try:
                        result = json.loads(response_text)
                        print(f"✅ Upload successful!")
                        print(f"📋 Response keys: {list(result.keys())}")
                        
                        # Check for expected response structure (Knowledge Engine format)
                        if result.get('status') == 'completed':
                            self.test_results['content_upload']['status'] = 'passed'
                            self.test_results['content_upload']['details'].append(f"Upload successful with {len(file_data)} bytes")
                            
                            # Store response for further testing
                            self.upload_response = result
                            
                            # Check for processing details
                            if 'chunks_created' in result:
                                chunks_count = result['chunks_created']
                                self.test_results['content_upload']['details'].append(f"Created {chunks_count} content chunks")
                            
                            if 'extracted_content_length' in result:
                                content_length = result['extracted_content_length']
                                self.test_results['content_upload']['details'].append(f"Extracted {content_length} characters of content")
                            
                            # Store job_id for further testing
                            self.job_id = result.get('job_id')
                                
                        else:
                            self.test_results['content_upload']['status'] = 'failed'
                            self.test_results['content_upload']['details'].append(f"Upload failed with status: {result.get('status', 'unknown')}")
                            
                    except json.JSONDecodeError as e:
                        print(f"❌ Invalid JSON response: {e}")
                        print(f"📄 Raw response: {response_text[:500]}...")
                        self.test_results['content_upload']['status'] = 'failed'
                        self.test_results['content_upload']['details'].append(f"Invalid JSON response: {str(e)}")
                        
                else:
                    print(f"❌ Upload failed with status {response.status}")
                    print(f"📄 Error response: {response_text[:500]}...")
                    self.test_results['content_upload']['status'] = 'failed'
                    self.test_results['content_upload']['details'].append(f"HTTP {response.status}: {response_text[:200]}")
                    
                    # Check for CORS issues
                    if response.status in [403, 405] or 'cors' in response_text.lower():
                        self.test_results['cors_issues']['status'] = 'failed'
                        self.test_results['cors_issues']['details'].append(f"CORS issue detected: {response.status}")
                    else:
                        self.test_results['cors_issues']['status'] = 'passed'
                        self.test_results['cors_issues']['details'].append("No CORS issues detected")
            
            # Cleanup test file
            try:
                os.remove(docx_path)
            except:
                pass
                
        except Exception as e:
            print(f"❌ Content upload test failed: {e}")
            self.test_results['content_upload']['status'] = 'failed'
            self.test_results['content_upload']['details'].append(f"Exception: {str(e)}")
    
    async def test_image_processing(self):
        """Test 2: Image Processing and Asset Library Integration"""
        print("\n🔍 TEST 2: Image Processing")
        
        try:
            # Check if we have upload response with processing data
            if not hasattr(self, 'upload_response'):
                self.test_results['image_processing']['status'] = 'skipped'
                self.test_results['image_processing']['details'].append('No upload response available')
                return
            
            upload_result = self.upload_response
            
            # For Knowledge Engine, check if content was processed
            content_length = upload_result.get('extracted_content_length', 0)
            print(f"📊 Content extracted: {content_length} characters")
            
            if content_length > 0:
                self.test_results['image_processing']['status'] = 'passed'
                self.test_results['image_processing']['details'].append(f"Successfully extracted {content_length} characters of content")
                
                # Test asset library endpoint to verify any images were saved
                assets_url = f"{API_BASE}/assets"
                
                try:
                    async with self.session.get(assets_url) as response:
                        if response.status == 200:
                            assets_data = await response.json()
                            if 'assets' in assets_data and assets_data['assets']:
                                recent_assets = [asset for asset in assets_data['assets'] 
                                               if asset.get('source') == 'docx_extraction']
                                
                                if recent_assets:
                                    self.test_results['image_processing']['details'].append(
                                        f"Found {len(recent_assets)} images in asset library"
                                    )
                                    
                                    # Test image URL accessibility
                                    for asset in recent_assets[:2]:  # Test first 2 images
                                        image_url = f"{BACKEND_URL}{asset['url']}"
                                        try:
                                            async with self.session.get(image_url) as img_response:
                                                if img_response.status == 200:
                                                    self.test_results['image_processing']['details'].append(
                                                        f"Image accessible: {asset['filename']}"
                                                    )
                                                else:
                                                    self.test_results['image_processing']['details'].append(
                                                        f"Image not accessible: {asset['filename']} (HTTP {img_response.status})"
                                                    )
                                        except Exception as img_error:
                                            self.test_results['image_processing']['details'].append(
                                                f"Image access error: {asset['filename']} - {str(img_error)}"
                                            )
                                else:
                                    self.test_results['image_processing']['details'].append(
                                        "No images from DOCX extraction found in asset library"
                                    )
                                
                        else:
                            self.test_results['image_processing']['details'].append(
                                f"Asset library not accessible (HTTP {response.status})"
                            )
                            
                except Exception as assets_error:
                    self.test_results['image_processing']['details'].append(
                        f"Asset library test failed: {str(assets_error)}"
                    )
                    
            else:
                self.test_results['image_processing']['status'] = 'failed'
                self.test_results['image_processing']['details'].append("No content was extracted from DOCX file")
                
        except Exception as e:
            print(f"❌ Image processing test failed: {e}")
            self.test_results['image_processing']['status'] = 'failed'
            self.test_results['image_processing']['details'].append(f"Exception: {str(e)}")
    
    async def test_content_coverage(self):
        """Test 3: Content Coverage and Content Library Integration"""
        print("\n🔍 TEST 3: Content Coverage")
        
        try:
            if not hasattr(self, 'upload_response'):
                self.test_results['content_coverage']['status'] = 'skipped'
                self.test_results['content_coverage']['details'].append('No upload response available')
                return
            
            upload_result = self.upload_response
            
            # Check content extraction results
            content_length = upload_result.get('extracted_content_length', 0)
            chunks_created = upload_result.get('chunks_created', 0)
            
            print(f"📊 Content extracted: {content_length} characters")
            print(f"📊 Chunks created: {chunks_created}")
            
            if content_length > 0 and chunks_created > 0:
                self.test_results['content_coverage']['status'] = 'passed'
                self.test_results['content_coverage']['details'].append(
                    f"Good content extraction: {content_length} characters, {chunks_created} chunks"
                )
                
                # Test Content Library to see if articles were created
                content_library_url = f"{API_BASE}/content-library"
                
                try:
                    async with self.session.get(content_library_url) as response:
                        if response.status == 200:
                            library_data = await response.json()
                            
                            if 'articles' in library_data and library_data['articles']:
                                articles = library_data['articles']
                                recent_articles = [article for article in articles 
                                                 if 'knowledge_engine_test' in article.get('title', '').lower() or
                                                    'knowledge_engine_test' in article.get('content', '').lower()]
                                
                                if recent_articles:
                                    self.test_results['content_coverage']['details'].append(
                                        f"Found {len(recent_articles)} articles in Content Library"
                                    )
                                    
                                    # Analyze article quality
                                    for i, article in enumerate(recent_articles[:3]):  # Check first 3
                                        title = article.get('title', 'Untitled')
                                        content = article.get('content', '')
                                        word_count = len(content.split()) if content else 0
                                        
                                        self.test_results['content_coverage']['details'].append(
                                            f"Article {i+1}: '{title[:50]}...' ({word_count} words)"
                                        )
                                        
                                        # Check for HTML structure
                                        if '<h1>' in content or '<h2>' in content:
                                            self.test_results['content_coverage']['details'].append(
                                                f"Article {i+1}: Has proper heading structure"
                                            )
                                        
                                        # Check for images
                                        if '<img' in content or '<figure' in content:
                                            self.test_results['content_coverage']['details'].append(
                                                f"Article {i+1}: Contains embedded images"
                                            )
                                else:
                                    self.test_results['content_coverage']['details'].append(
                                        "No articles from test document found in Content Library"
                                    )
                            else:
                                self.test_results['content_coverage']['details'].append(
                                    "Content Library is empty or not accessible"
                                )
                                
                        else:
                            self.test_results['content_coverage']['details'].append(
                                f"Content Library not accessible (HTTP {response.status})"
                            )
                            
                except Exception as library_error:
                    self.test_results['content_coverage']['details'].append(
                        f"Content Library test failed: {str(library_error)}"
                    )
                    
            elif content_length > 0:
                self.test_results['content_coverage']['status'] = 'partial'
                self.test_results['content_coverage']['details'].append(
                    f"Content extracted ({content_length} chars) but no chunks created"
                )
            else:
                self.test_results['content_coverage']['status'] = 'failed'
                self.test_results['content_coverage']['details'].append('No content was extracted from the document')
                
        except Exception as e:
            print(f"❌ Content coverage test failed: {e}")
            self.test_results['content_coverage']['status'] = 'failed'
            self.test_results['content_coverage']['details'].append(f"Exception: {str(e)}")
    
    async def test_article_generation_quality(self):
        """Test 4: Article Generation and Image Placement Quality"""
        print("\n🔍 TEST 4: Article Generation Quality")
        
        try:
            if not hasattr(self, 'upload_response'):
                self.test_results['article_generation']['status'] = 'skipped'
                self.test_results['article_generation']['details'].append('No upload response available')
                return
            
            upload_result = self.upload_response
            articles = upload_result.get('articles', [])
            
            if not articles:
                self.test_results['article_generation']['status'] = 'failed'
                self.test_results['article_generation']['details'].append('No articles to analyze')
                return
            
            quality_score = 0
            max_score = 0
            
            for i, article in enumerate(articles):
                if isinstance(article, dict):
                    content = article.get('content', '') or article.get('html', '')
                    title = article.get('title', '')
                    
                    # Test 1: Has meaningful title
                    max_score += 1
                    if title and len(title.strip()) > 5:
                        quality_score += 1
                        self.test_results['article_generation']['details'].append(
                            f"Article {i+1}: Good title - '{title[:50]}...'"
                        )
                    
                    # Test 2: Has substantial content
                    max_score += 1
                    if content and len(content.strip()) > 200:
                        quality_score += 1
                        self.test_results['article_generation']['details'].append(
                            f"Article {i+1}: Substantial content ({len(content)} chars)"
                        )
                    
                    # Test 3: Proper HTML structure
                    max_score += 1
                    if '<h1>' in content or '<h2>' in content:
                        quality_score += 1
                        self.test_results['article_generation']['details'].append(
                            f"Article {i+1}: Proper HTML heading structure"
                        )
                    
                    # Test 4: Image integration
                    max_score += 1
                    if '<img' in content or '<figure' in content:
                        quality_score += 1
                        
                        # Count images in this article
                        img_count = content.count('<img')
                        figure_count = content.count('<figure')
                        
                        self.test_results['article_generation']['details'].append(
                            f"Article {i+1}: Contains {img_count} images, {figure_count} figures"
                        )
                    
                    # Test 5: Professional formatting
                    max_score += 1
                    if '<p>' in content and ('</p>' in content):
                        quality_score += 1
                        self.test_results['article_generation']['details'].append(
                            f"Article {i+1}: Professional paragraph formatting"
                        )
            
            # Calculate overall quality score
            if max_score > 0:
                quality_percentage = (quality_score / max_score) * 100
                
                if quality_percentage >= 80:
                    self.test_results['article_generation']['status'] = 'passed'
                    self.test_results['article_generation']['details'].append(
                        f"Excellent quality: {quality_percentage:.1f}% ({quality_score}/{max_score})"
                    )
                elif quality_percentage >= 60:
                    self.test_results['article_generation']['status'] = 'passed'
                    self.test_results['article_generation']['details'].append(
                        f"Good quality: {quality_percentage:.1f}% ({quality_score}/{max_score})"
                    )
                else:
                    self.test_results['article_generation']['status'] = 'failed'
                    self.test_results['article_generation']['details'].append(
                        f"Poor quality: {quality_percentage:.1f}% ({quality_score}/{max_score})"
                    )
            else:
                self.test_results['article_generation']['status'] = 'failed'
                self.test_results['article_generation']['details'].append('No quality metrics available')
                
        except Exception as e:
            print(f"❌ Article generation test failed: {e}")
            self.test_results['article_generation']['status'] = 'failed'
            self.test_results['article_generation']['details'].append(f"Exception: {str(e)}")
    
    async def run_all_tests(self):
        """Run all Knowledge Engine tests"""
        print("🚀 Starting Knowledge Engine Testing Suite")
        print(f"🔗 Backend URL: {BACKEND_URL}")
        print(f"🔗 API Base: {API_BASE}")
        
        await self.setup_session()
        
        try:
            # Run tests in sequence
            await self.test_content_upload_api()
            await self.test_image_processing()
            await self.test_content_coverage()
            await self.test_article_generation_quality()
            
        finally:
            await self.cleanup_session()
        
        # Print results summary
        self.print_test_summary()
    
    def print_test_summary(self):
        """Print comprehensive test results"""
        print("\n" + "="*80)
        print("🎯 KNOWLEDGE ENGINE TEST RESULTS SUMMARY")
        print("="*80)
        
        passed_tests = 0
        total_tests = 0
        
        for test_name, result in self.test_results.items():
            total_tests += 1
            status = result['status']
            
            if status == 'passed':
                status_icon = "✅"
                passed_tests += 1
            elif status == 'failed':
                status_icon = "❌"
            elif status == 'skipped':
                status_icon = "⏭️"
            else:
                status_icon = "⏳"
            
            print(f"\n{status_icon} {test_name.upper().replace('_', ' ')}: {status.upper()}")
            
            for detail in result['details']:
                print(f"   • {detail}")
        
        print(f"\n📊 OVERALL RESULTS: {passed_tests}/{total_tests} tests passed")
        
        if passed_tests == total_tests:
            print("🎉 ALL TESTS PASSED - Knowledge Engine is fully operational!")
        elif passed_tests >= total_tests * 0.75:
            print("✅ MOSTLY WORKING - Knowledge Engine has good functionality with minor issues")
        elif passed_tests >= total_tests * 0.5:
            print("⚠️ PARTIALLY WORKING - Knowledge Engine has significant issues that need attention")
        else:
            print("❌ MAJOR ISSUES - Knowledge Engine requires substantial fixes")
        
        print("="*80)

async def main():
    """Main test execution"""
    tester = KnowledgeEngineTest()
    await tester.run_all_tests()

if __name__ == "__main__":
    asyncio.run(main())